import json
import os
import re
from datetime import datetime, timedelta

import pandas as pd
import pythoncom
from docx import Document
from dotenv import load_dotenv
from win32com.client import Dispatch

from src.check_volume import validate_dataframe
from main import parse_args


def extract_route_info(parags, add_number):
    for i, para in enumerate(parags):
        if i == 0:
            sourse = para.text.split("\n")[0] + f" ДС{add_number}"
        if para.text.startswith("Параметры маршрута"):
            route = para.text.split()[-1]
            date_range = parags[i + 1].text.split("\n")[0]
            return sourse, route, date_range
    return None, None, None

def extract_km_data(parags, string1="Протяжённость, всего", string2="Схема маршрута."):
    text_between = []
    found_string1 = False
    for para in parags:
        if string1 in para.text:
            match = re.search(r"\d+,\d+", para.text)
            text_between.append(match.group(0) if match else "0")
            found_string1 = True
        elif string2 in para.text:
            break
        elif found_string1:
            match = re.search(r"\d+,\d+", para.text)
            if match or para.text.strip():
                text_between.append(match.group(0) if match else "0")
    return text_between[:3] if len(text_between) >= 3 else ["0", "0", "0"]

def parse_table(table, filename, folder_name, forward_km, reverse_km):
    df = [[cell.text for cell in row.cells] for row in table.rows]
    flights_table = pd.DataFrame(df)

    new_header = pd.MultiIndex.from_arrays(
        [flights_table.iloc[1], flights_table.iloc[0]]
    )
    flights_table.columns = new_header
    flights_table = flights_table[2:]
    flights_table.set_index(flights_table.columns[0], inplace=True)

    validate_dataframe(filename, folder_name, flights_table, forward_km, reverse_km)

    flights_table = flights_table.loc[
        :, flights_table.columns.get_level_values(1).str.contains("Количество рейсов")
    ]
    flights_table.columns = flights_table.columns.get_level_values(0)
    flights_table = flights_table.rename(index={"Прямое": "Прямо", "Обратное": "Обратно"})
    flights_table = flights_table.iloc[:-1]
    return prepare_flights(flights_table.to_dict())

def exctract_data_from_app(path, agg_num, add_number, filename):
    document = Document(path)
    parags = document.paragraphs
    table_count = len(document.tables)

    sourse, route, date_range = extract_route_info(parags, add_number)
    total_km, forward_km, reverse_km = extract_km_data(parags)

    folder_name = path.split("\\")[-2]

    if table_count in [3, 4, 5]:
        fl_by_days = parse_table(document.tables[-1], filename, folder_name, forward_km, reverse_km)
        dt_dict = create_date_dict(date_range)
        return pd.DataFrame([[route, dt_dict, total_km, forward_km, reverse_km, fl_by_days, sourse, agg_num]])

    elif table_count == 6:
        win_sum = [[], []]
        counter = 0
        for i in range(2):
            win_sum[i].extend([route, [create_date_dict(date_range)]])
        for para in parags:
            if "Количество рейсов и пробег транспортных средств" in para.text:
                dt_dict_season = create_date_dict(para.text)
                win_sum[counter][1].append(dt_dict_season)
                counter += 1

        for idx, table in enumerate(document.tables[-2:]):
            fl_by_days = parse_table(table, filename, folder_name, forward_km, reverse_km)
            win_sum[idx].extend([total_km, forward_km, reverse_km, fl_by_days, sourse, agg_num])
        return pd.DataFrame(win_sum)

    else:
        print(f"странное кол-во таблиц в файле {path}, {table_count}")
        return 0


def convert_doc_to_docx(doc_path):
    try:
        # Проверяем, что путь указывает на файл с расширением .doc
        if not doc_path.lower().endswith(".doc"):
            raise ValueError("The file must have a .doc extension")

        # Инициализируем Word
        pythoncom.CoInitialize()
        word = Dispatch("Word.Application")
        word.Visible = False

        # Открываем документ
        doc = word.Documents.Open(doc_path)

        # Генерируем путь для сохранения нового документа в формате .docx
        docx_path = os.path.splitext(doc_path)[0] + ".docx"

        # Сохраняем документ в формате .docx
        doc.SaveAs(docx_path, FileFormat=16)  # 16 - формат для .docx

        # Закрываем документ без сохранения изменений
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()

        return docx_path
    except Exception as e:
        print(f"Ошибка: {e}")
        return None


def create_date_dict(input_string, lable_date=None):

    if "на" in input_string:
        r = input_string
        input_string = (
            r.replace("на", "начало_интервала")
            + " "
            + input_string.replace("на", "конец_интервала")
        )
    elif "по" in input_string:
        input_string = input_string.replace("с", "начало_интервала")
        input_string = input_string.replace("по", "конец_интервала")
    else:
        input_string = input_string.replace("с", "начало")

    date_pattern = r"(\b\w+\b)\s(\b\d{1,2}\.\d{1,2}(?:\.\d{2,4})?\b)"
    matches = re.findall(date_pattern, input_string)

    date_dict = {}
    for match in matches:
        word, date_str = match
        try:

            if len(date_str.split(".")) == 3 and len(date_str.split(".")[-1]) == 2:
                date = datetime.strptime(date_str, "%d.%m.%y")
                date_str = date.strftime("%d.%m.%Y")  # Преобразуем в полный формат
            else:
                date = datetime.strptime(date_str, "%d.%m.%Y")

            date_dict[word.lower()] = date.strftime("%d.%m.%Y")
        except ValueError:
            print("Неверный формат даты")

    return date_dict


def prepare_flights(my_dict):
    weekdays_dict = {
        "Рабочие дни": [1, 2, 3, 4, 5],
        "Выходные и праздничные дни": [6, 7],
        "Рабочие дни кроме пятницы": [1, 2, 3, 4],
        "Пятница": [5],
        "Субботние дни": [6],
        "Воскресные и праздничные дни": [7],
        "Рабочие, выходные и праздничные дни": [1, 2, 3, 4, 5, 6, 7],
        "Рабочие и субботние дни": [1, 2, 3, 4, 5, 6],
        "Пятница, выходные и праздничные дни": [5, 6, 7],
    }
    result_dict = {}
    for k, v in my_dict.items():
        for weekday in weekdays_dict[k]:
            for key, value in my_dict[k].items():
                result_dict[f"{key} - {weekday}"] = value

    return result_dict


def str_to_dict(cell):
    try:
        # Используем json.loads для преобразования строки в словарь
        return json.loads(cell)
    except (ValueError, TypeError):
        # Если преобразование не удалось, возвращаем ячейку как есть
        return cell


def expand_dates(row):
    columns = [
        "Маршрут",
        "Дата",
        "Всего",
        "от НП",
        "от КП",
        "Кол-во рейсов",
        "Источник",
        "Номер ГК",
        "Тип даты",
    ]
    rows = pd.DataFrame(columns=columns)

    dates = row["Дата"]
    if isinstance(dates, dict):
        dates = [dates]

    for date_dict in dates:
        for key, value in date_dict.items():
            if key in [
                "начало",
                "начало_интервала",
                "точка",
            ]:
                new_row = row.copy()
                new_row["Тип даты"] = key
                new_row["Дата"] = value
                rows = pd.concat([rows, pd.DataFrame([new_row])], ignore_index=True)
            if key == "конец_интервала":
                empty_row = row.copy()
                empty_row[:] = ""
                empty_row["Маршрут"] = row["Маршрут"]
                empty_row["Тип даты"] = key
                date = datetime.strptime(value, "%d.%m.%Y") + timedelta(days=1)
                empty_row["Дата"] = date.strftime("%d.%m.%Y")
                empty_row["Источник"] = row["Источник"]
                empty_row["Номер ГК"] = row["Номер ГК"]
                rows = pd.concat([rows, pd.DataFrame([empty_row])], ignore_index=True)
    return rows


def main():
    args = parse_args()
    AGG_NUM = args.AGG_NUMS[0]
    ADD_NUMBER = args.ADD_NUMBERS[0]
    load_dotenv()

    filename = os.path.join(
        os.environ["path_to_results"],
        f"проверка объемов в приложениях ДС{AGG_NUM} ГК{ADD_NUMBER}.txt",
    )
    apps = os.environ["path_to_apps"]
    apps = apps.replace("AGG_NUM", str(AGG_NUM))
    apps = apps.replace("ADD_NUMBER", str(ADD_NUMBER))

    directories = [d for d in os.listdir(apps) if os.path.isdir(os.path.join(apps, d))]

    result_df = pd.DataFrame()
    for directory in directories:
        d = directory
        print(f"Обрабатываем папку: {directory}")
        os.chdir(apps + "/" + directory)  # Переходим в текущую папку
        files = os.listdir("..")  # Получаем список файлов в текущей папке

        doc_and_docx_files = [
            file
            for file in files
            if (file.lower().endswith(".doc") or file.lower().endswith(".docx"))
            and not file.startswith("~")
        ]

        if len(doc_and_docx_files) == 0:
            print(f"В папке {directory} нет файлов с расширением .doc или .docx")
        else:
            for file in doc_and_docx_files:
                if file.lower().endswith(".docx"):
                    docx_path = os.path.abspath(file)
                    data_from_app = exctract_data_from_app(
                        docx_path, AGG_NUM, ADD_NUMBER, filename
                    )
                    result_df = pd.concat(
                        [result_df, data_from_app], ignore_index=True, axis=0
                    )
                    break
                if file.lower().endswith(".doc"):
                    doc_path = os.path.abspath(file)
                    docx_path = convert_doc_to_docx(doc_path)
                    # os.remove(file)
                    if docx_path:
                        data_from_app = exctract_data_from_app(
                            docx_path, AGG_NUM, ADD_NUMBER, filename
                        )
                        result_df = pd.concat(
                            [result_df, pd.DataFrame(data_from_app)],
                            ignore_index=True,
                            axis=0,
                        )

                    else:
                        print(f"Не удалось конвертировать файл {file}")

        os.chdir("../..")
    columns = [
        "Маршрут",
        "Дата",
        "Всего",
        "от НП",
        "от КП",
        "Кол-во рейсов",
        "Источник",
        "Номер ГК",
    ]
    result_df.columns = columns
    result_df["Тип даты"] = ""

    result_df["Дата"] = result_df["Дата"].apply(str_to_dict)
    result_df["Кол-во рейсов"] = result_df["Кол-во рейсов"].apply(str_to_dict)

    exp_rows = result_df.apply(lambda row: expand_dates(row), axis=1)
    result_df = pd.concat(exp_rows.values, ignore_index=True)

    df_flights = result_df["Кол-во рейсов"].apply(pd.Series)
    result_df = pd.concat(
        [result_df.drop(columns=["Кол-во рейсов"]), df_flights], axis=1
    ).drop_duplicates()

    duplicates = result_df[result_df.duplicated(subset=["Маршрут", "Дата"], keep=False)]
    duplicates = duplicates[duplicates["Тип даты"] == "конец_интервала"]
    indices_to_delete = duplicates.index
    res = result_df.drop(indices_to_delete).reset_index(drop=True)

    res.to_excel(
        os.path.join(
            os.environ["path_to_results"], f"приложения {AGG_NUM} {ADD_NUMBER}.xlsx"
        ),
        index=False,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
